import sys
import subprocess

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
from docx import Document

md_file = r"e:\Project\Class\OSG\Model_Evaluation_Report.md"
docx_file = r"e:\Project\Class\OSG\Model_Evaluation_Report.docx"

doc = Document()
doc.add_heading('BÁO CÁO ĐÁNH GIÁ MÔ HÌNH DỰ BÁO TRÀN RAM', 0)

with open(md_file, 'r', encoding='utf-8') as file:
    lines = file.readlines()

for line in lines:
    line = line.strip()
    if not line or line.startswith('# BÁO CÁO') or line == "---":
        continue
    
    if line.startswith('### '):
        doc.add_heading(line[4:].replace('*',''), level=3)
    elif line.startswith('## '):
        doc.add_heading(line[3:].replace('*',''), level=2)
    elif line.startswith('# '):
        doc.add_heading(line[2:].replace('*',''), level=1)
    elif line.startswith('* '):
        doc.add_paragraph(line[2:].replace('**', ''), style='List Bullet')
    elif line.startswith('- '):
        doc.add_paragraph(line[2:].replace('**', ''), style='List Bullet')
    else:
        # Xóa các markdown đậm (**)
        clean_text = line.replace('**', '')
        doc.add_paragraph(clean_text)

doc.save(docx_file)
print(f"Đã chuyển đổi thành công sang file: {docx_file}")
